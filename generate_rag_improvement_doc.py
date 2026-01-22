#!/usr/bin/env python3
"""
Script to generate a Word document with phases for improving document RAG system
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime

def add_heading_with_style(doc, text, level=1):
    """Add a heading with custom styling"""
    heading = doc.add_heading(text, level=level)
    heading.runs[0].font.color.rgb = RGBColor(0, 51, 102)
    return heading

def add_phase(doc, phase_number, title, description, activities):
    """Add a phase section to the document"""
    # Phase title
    heading = add_heading_with_style(doc, f"Fase {phase_number}: {title}", level=2)
    
    # Description
    desc_para = doc.add_paragraph()
    desc_para.add_run("Descrizione: ").bold = True
    desc_para.add_run(description)
    
    # Activities
    doc.add_paragraph("Attività principali:", style='Heading 3')
    for activity in activities:
        p = doc.add_paragraph(activity, style='List Bullet')
        p.paragraph_format.left_indent = Inches(0.5)
    
    doc.add_paragraph()  # Add spacing

def add_metrics_section(doc, metrics):
    """Add a metrics section"""
    doc.add_paragraph("Metriche di successo:", style='Heading 3')
    for metric in metrics:
        p = doc.add_paragraph(metric, style='List Bullet')
        p.paragraph_format.left_indent = Inches(0.5)
    
    doc.add_paragraph()

def create_rag_improvement_document():
    """Create the main document"""
    doc = Document()
    
    # Title
    title = doc.add_heading('Piano di Miglioramento RAG Documentale', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Subtitle with date
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(f'Sistema DocN - {datetime.now().strftime("%B %Y")}')
    run.italic = True
    run.font.size = Pt(12)
    
    doc.add_paragraph()
    
    # Executive Summary
    add_heading_with_style(doc, '📋 Sommario Esecutivo', level=1)
    summary_text = (
        "Questo documento delinea un approccio strutturato in fasi per migliorare "
        "progressivamente il sistema RAG (Retrieval-Augmented Generation) documentale. "
        "L'obiettivo è ottimizzare la qualità delle risposte, ridurre le allucinazioni, "
        "migliorare la rilevanza dei risultati e garantire un'esperienza utente eccellente."
    )
    doc.add_paragraph(summary_text)
    doc.add_paragraph()
    
    # Current State Analysis
    add_heading_with_style(doc, '🔍 Analisi dello Stato Attuale', level=1)
    
    current_features = [
        "Embedding vettoriali per ricerca semantica",
        "Sistema di caching per ottimizzazione prestazioni",
        "Query analysis con HyDE e query rewriting",
        "Reranking con considerazione di diversità (MMR)",
        "Metriche di qualità RAG con RAGAS",
        "Rilevamento di allucinazioni",
        "Verifica delle citazioni",
        "Configurazione modulare per ogni fase RAG"
    ]
    
    doc.add_paragraph("Il sistema attuale include:", style='Heading 3')
    for feature in current_features:
        doc.add_paragraph(f"✓ {feature}", style='List Bullet')
    
    doc.add_paragraph()
    
    # Improvement Phases
    add_heading_with_style(doc, '🎯 Fasi di Miglioramento', level=1)
    
    # Phase 1
    add_phase(
        doc,
        1,
        "Ottimizzazione degli Embeddings",
        "Migliorare la qualità della rappresentazione vettoriale dei documenti per una ricerca semantica più accurata.",
        [
            "Valutare modelli di embedding alternativi (es. Multilingual-E5, BGE-M3)",
            "Implementare fine-tuning del modello di embedding su documenti specifici del dominio",
            "Ottimizzare il chunking dei documenti (dimensione, overlap, strategie smart)",
            "Implementare chunking semantico basato su struttura del documento",
            "Aggiungere metadati ricchi per ogni chunk (titolo, sezione, keywords)",
            "Validare miglioramenti con metriche di retrieval (MRR, NDCG)"
        ]
    )
    
    add_metrics_section(doc, [
        "Mean Reciprocal Rank (MRR) > 0.85",
        "Normalized Discounted Cumulative Gain (NDCG) > 0.80",
        "Riduzione del 30% dei casi di recupero non rilevante"
    ])
    
    # Phase 2
    add_phase(
        doc,
        2,
        "Miglioramento della Ricerca Ibrida",
        "Combinare efficacemente ricerca vettoriale e ricerca testuale per massimizzare la rilevanza.",
        [
            "Implementare ricerca ibrida con BM25 + Vector Search",
            "Ottimizzare il peso relativo tra ricerca semantica e keyword-based",
            "Aggiungere filtri avanzati (data, tipo documento, autore)",
            "Implementare ricerca multi-hop per query complesse",
            "Abilitare query expansion con sinonimi e termini correlati",
            "Implementare cache semantica per query simili"
        ]
    )
    
    add_metrics_section(doc, [
        "Precision@5 > 0.90",
        "Recall@10 > 0.85",
        "Tempo medio di risposta < 1.5 secondi"
    ])
    
    # Phase 3
    add_phase(
        doc,
        3,
        "Reranking Avanzato",
        "Riordinare i risultati recuperati utilizzando modelli più sofisticati per massimizzare la rilevanza.",
        [
            "Integrare modelli di reranking specializzati (es. Cohere Rerank, bge-reranker)",
            "Implementare reranking cross-encoder per precisione maggiore",
            "Aggiungere pesi temporali per privilegiare documenti recenti",
            "Ottimizzare il parametro MMR Lambda per bilanciamento rilevanza/diversità",
            "Implementare contextual reranking basato su conversazione",
            "A/B testing di diverse strategie di reranking"
        ]
    )
    
    add_metrics_section(doc, [
        "Miglioramento del 25% nella rilevanza percepita dagli utenti",
        "Riduzione del 40% dei click su risultati non rilevanti",
        "Aumento del 30% del tempo di permanenza sulle risposte"
    ])
    
    # Phase 4
    add_phase(
        doc,
        4,
        "Generazione e Sintesi Avanzata",
        "Migliorare la qualità della generazione della risposta finale utilizzando tecniche avanzate di prompt engineering e LLM.",
        [
            "Ottimizzare i prompt di sistema per risposte più accurate e contestuali",
            "Implementare chain-of-thought reasoning per query complesse",
            "Abilitare compressione contestuale per gestire più informazioni",
            "Implementare self-consistency checking (multiple generations + voting)",
            "Aggiungere fact-checking automatico delle risposte generate",
            "Implementare refinement iterativo per risposte di alta qualità",
            "Gestire meglio le citazioni con link diretti ai documenti fonte"
        ]
    )
    
    add_metrics_section(doc, [
        "Faithfulness score (RAGAS) > 0.90",
        "Answer relevancy score (RAGAS) > 0.85",
        "Riduzione del 50% delle allucinazioni rilevate"
    ])
    
    # Phase 5
    add_phase(
        doc,
        5,
        "Monitoraggio e Quality Assurance",
        "Implementare un sistema robusto di monitoraggio continuo e miglioramento della qualità.",
        [
            "Dashboard real-time per metriche RAG (latency, relevance, quality)",
            "Sistema di logging completo per analisi post-mortem",
            "Implementare human-in-the-loop feedback per miglioramento continuo",
            "A/B testing framework per confrontare miglioramenti",
            "Alert automatici per degradazione della qualità",
            "Raccolta e analisi di query fallite per training futuro",
            "Implementare regression testing per prevenire degradazioni"
        ]
    )
    
    add_metrics_section(doc, [
        "Copertura di monitoring su 100% delle query",
        "Tempo medio di detection dei problemi < 5 minuti",
        "Feedback degli utenti raccolto su almeno il 10% delle query"
    ])
    
    # Phase 6
    add_phase(
        doc,
        6,
        "Ottimizzazione Prestazioni e Scalabilità",
        "Garantire che il sistema RAG possa gestire un carico crescente mantenendo performance eccellenti.",
        [
            "Implementare caching multi-livello (query, embeddings, risultati)",
            "Ottimizzare le query al database vettoriale (batch processing, indexing)",
            "Implementare pre-computation degli embeddings per documenti statici",
            "Configurare auto-scaling basato sul carico",
            "Ottimizzare l'utilizzo della memoria e delle risorse GPU/CPU",
            "Implementare rate limiting e throttling intelligente",
            "Load testing e stress testing del sistema"
        ]
    )
    
    add_metrics_section(doc, [
        "Throughput > 100 query/secondo",
        "P95 latency < 2 secondi",
        "Utilizzo risorse ottimale (CPU < 70%, memoria < 80%)"
    ])
    
    # Phase 7
    add_phase(
        doc,
        7,
        "Personalizzazione e Context Awareness",
        "Rendere il sistema più intelligente adattandolo al contesto e alle preferenze dell'utente.",
        [
            "Implementare user profiling per personalizzare i risultati",
            "Context tracking per conversazioni multi-turn più coerenti",
            "Apprendimento dalle interazioni utente (implicit feedback)",
            "Personalizzazione del linguaggio e del tono delle risposte",
            "Supporto multi-lingua con modelli specifici per lingua",
            "Memorizzazione delle preferenze dell'utente",
            "Raccomandazioni proattive basate su query precedenti"
        ]
    )
    
    add_metrics_section(doc, [
        "Aumento del 40% della soddisfazione utente",
        "Riduzione del 30% delle query di follow-up",
        "Aumento del 25% del tasso di utilizzo ripetuto"
    ])
    
    # Implementation Timeline
    add_heading_with_style(doc, '📅 Timeline di Implementazione Suggerita', level=1)
    
    timeline_table = doc.add_table(rows=8, cols=3)
    timeline_table.style = 'Light Grid Accent 1'
    
    # Header
    header_cells = timeline_table.rows[0].cells
    header_cells[0].text = 'Fase'
    header_cells[1].text = 'Durata Stimata'
    header_cells[2].text = 'Priorità'
    
    # Make header bold
    for cell in header_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
    
    # Data
    phases_data = [
        ('Fase 1: Ottimizzazione Embeddings', '2-3 settimane', 'Alta'),
        ('Fase 2: Ricerca Ibrida', '3-4 settimane', 'Alta'),
        ('Fase 3: Reranking Avanzato', '2-3 settimane', 'Media'),
        ('Fase 4: Generazione Avanzata', '3-4 settimane', 'Alta'),
        ('Fase 5: Monitoraggio e QA', '2-3 settimane', 'Alta'),
        ('Fase 6: Prestazioni e Scalabilità', '2-3 settimane', 'Media'),
        ('Fase 7: Personalizzazione', '4-5 settimane', 'Bassa'),
    ]
    
    for i, (phase, duration, priority) in enumerate(phases_data, 1):
        row_cells = timeline_table.rows[i].cells
        row_cells[0].text = phase
        row_cells[1].text = duration
        row_cells[2].text = priority
    
    doc.add_paragraph()
    
    # Best Practices
    add_heading_with_style(doc, '💡 Best Practices Generali', level=1)
    
    best_practices = [
        "Implementare ogni fase in modo incrementale con A/B testing",
        "Validare ogni miglioramento con metriche oggettive prima di procedere",
        "Mantenere sempre una versione stabile in produzione (rollback ready)",
        "Documentare tutte le modifiche e i risultati dei test",
        "Coinvolgere gli utenti finali per feedback qualitativo",
        "Automatizzare il testing per prevenire regressioni",
        "Monitorare costantemente costi e performance",
        "Rivedere periodicamente le priorità in base ai risultati"
    ]
    
    for practice in best_practices:
        p = doc.add_paragraph(f"• {practice}")
        p.paragraph_format.left_indent = Inches(0.25)
    
    doc.add_paragraph()
    
    # Tools and Technologies
    add_heading_with_style(doc, '🛠️ Strumenti e Tecnologie Consigliate', level=1)
    
    doc.add_paragraph("Embedding Models:", style='Heading 3')
    embedding_tools = [
        "Multilingual-E5-large (per documenti multilingua)",
        "BGE-M3 (per supporto multilingua robusto)",
        "OpenAI text-embedding-3-large (per qualità superiore)"
    ]
    for tool in embedding_tools:
        doc.add_paragraph(tool, style='List Bullet')
    
    doc.add_paragraph("Reranking:", style='Heading 3')
    reranking_tools = [
        "Cohere Rerank API",
        "bge-reranker-large",
        "Cross-encoders (sentence-transformers)"
    ]
    for tool in reranking_tools:
        doc.add_paragraph(tool, style='List Bullet')
    
    doc.add_paragraph("Vector Databases:", style='Heading 3')
    vector_db_tools = [
        "Qdrant (attuale, ottimizzare configurazione)",
        "Pinecone (per scalabilità cloud)",
        "Weaviate (per ricerca ibrida avanzata)"
    ]
    for tool in vector_db_tools:
        doc.add_paragraph(tool, style='List Bullet')
    
    doc.add_paragraph("Monitoring:", style='Heading 3')
    monitoring_tools = [
        "LangSmith (per tracing completo LLM)",
        "Prometheus + Grafana (per metriche real-time)",
        "Sentry (per error tracking)",
        "Custom dashboard con Application Insights"
    ]
    for tool in monitoring_tools:
        doc.add_paragraph(tool, style='List Bullet')
    
    doc.add_paragraph()
    
    # Conclusion
    add_heading_with_style(doc, '🎓 Conclusioni', level=1)
    
    conclusion_text = (
        "Il miglioramento di un sistema RAG documentale è un processo iterativo che richiede "
        "attenzione continua e validazione empirica. Le fasi proposte in questo documento "
        "forniscono una roadmap strutturata per evolvere progressivamente il sistema, "
        "bilanciando qualità delle risposte, performance, e esperienza utente.\n\n"
        "È fondamentale procedere in modo incrementale, validando ogni miglioramento con "
        "metriche oggettive e feedback degli utenti prima di passare alla fase successiva. "
        "Il successo a lungo termine dipenderà dalla capacità di monitorare costantemente "
        "il sistema e adattarsi alle esigenze emergenti degli utenti."
    )
    doc.add_paragraph(conclusion_text)
    
    doc.add_paragraph()
    
    # Footer
    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_para.add_run(
        f"Documento generato il {datetime.now().strftime('%d/%m/%Y alle %H:%M')}\n"
        "Sistema DocN - Archiviazione Documentale con RAG"
    )
    footer_run.italic = True
    footer_run.font.size = Pt(9)
    footer_run.font.color.rgb = RGBColor(128, 128, 128)
    
    return doc

def main():
    """Main execution"""
    print("Generazione del documento di miglioramento RAG...")
    
    doc = create_rag_improvement_document()
    
    output_filename = "PIANO_MIGLIORAMENTO_RAG.docx"
    doc.save(output_filename)
    
    print(f"✓ Documento creato con successo: {output_filename}")
    print(f"  Il documento contiene 7 fasi dettagliate per migliorare il sistema RAG")

if __name__ == "__main__":
    main()
